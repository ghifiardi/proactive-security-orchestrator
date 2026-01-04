#!/usr/bin/env python3
"""Convert markdown documentation to DOCX format."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def add_styled_paragraph(doc, text, style='Normal'):
    """Add a paragraph with specific style."""
    p = doc.add_paragraph(text, style=style)
    return p


def add_heading_with_formatting(doc, text, level=1):
    """Add a heading with custom formatting."""
    heading = doc.add_heading(text, level=level)
    return heading


def add_code_block(doc, code_text):
    """Add a code block with monospace font and gray background."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Set paragraph background color (light gray)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table_from_markdown(doc, table_lines):
    """Parse markdown table and create docx table."""
    # Parse table headers and rows
    rows_data = []
    for line in table_lines:
        if '|' in line and not all(c in '|-: ' for c in line):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows_data.append(cells)

    if not rows_data or len(rows_data) < 2:
        return None

    # Create table
    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Populate table
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                # Bold header row
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

    return table


def parse_markdown_line(doc, line, in_code_block, code_lines, in_table, table_lines):
    """Parse a single line of markdown and add to document."""
    # Handle code blocks
    if line.strip().startswith('```'):
        if in_code_block:
            # End of code block
            if code_lines:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines.clear()
            return False, code_lines, in_table, table_lines
        else:
            # Start of code block
            return True, code_lines, in_table, table_lines

    if in_code_block:
        code_lines.append(line.rstrip())
        return in_code_block, code_lines, in_table, table_lines

    # Handle tables
    if line.strip().startswith('|'):
        if not in_table:
            table_lines = [line]
            return in_code_block, code_lines, True, table_lines
        else:
            table_lines.append(line)
            return in_code_block, code_lines, in_table, table_lines
    elif in_table:
        # End of table
        if table_lines:
            add_table_from_markdown(doc, table_lines)
            table_lines = []
        in_table = False

    # Handle headings
    if line.startswith('#'):
        match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if match:
            level = len(match.group(1))
            text = match.group(2).strip()
            add_heading_with_formatting(doc, text, level=min(level, 3))
            return in_code_block, code_lines, in_table, table_lines

    # Handle horizontal rules
    if line.strip() in ['---', '***', '___']:
        doc.add_paragraph('_' * 80)
        return in_code_block, code_lines, in_table, table_lines

    # Handle bullet lists
    if line.strip().startswith(('- ', '* ', '+ ')):
        text = line.strip()[2:]
        p = doc.add_paragraph(text, style='List Bullet')
        return in_code_block, code_lines, in_table, table_lines

    # Handle numbered lists
    if re.match(r'^\d+\.\s+', line.strip()):
        text = re.sub(r'^\d+\.\s+', '', line.strip())
        p = doc.add_paragraph(text, style='List Number')
        return in_code_block, code_lines, in_table, table_lines

    # Handle checkboxes
    if '- [x]' in line or '- [ ]' in line:
        text = line.replace('- [x]', '☑').replace('- [ ]', '☐').strip()
        p = doc.add_paragraph(text, style='List Bullet')
        return in_code_block, code_lines, in_table, table_lines

    # Handle bold, italic, inline code
    if line.strip():
        # Process markdown formatting
        text = line
        # Remove excessive bold markers
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'__(.+?)__', r'\1', text)
        # Remove inline code markers
        text = re.sub(r'`(.+?)`', r'\1', text)
        # Clean up emoji/special characters
        text = text.replace('✅', '[OK]').replace('⚠️', '[WARN]').replace('🔴', '[CRITICAL]')
        text = text.replace('🟡', '[WARN]').replace('🟢', '[OK]')

        if text.strip():
            p = doc.add_paragraph(text.strip())
            return in_code_block, code_lines, in_table, table_lines

    return in_code_block, code_lines, in_table, table_lines


def convert_markdown_to_docx(markdown_file, docx_file):
    """Convert markdown file to DOCX format."""
    print(f"Converting {markdown_file} to {docx_file}...")

    # Create document
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Read markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Parse markdown
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    for line in lines:
        in_code_block, code_lines, in_table, table_lines = parse_markdown_line(
            doc, line.rstrip(), in_code_block, code_lines, in_table, table_lines
        )

    # Handle any remaining code block
    if code_lines:
        add_code_block(doc, '\n'.join(code_lines))

    # Handle any remaining table
    if table_lines:
        add_table_from_markdown(doc, table_lines)

    # Save document
    doc.save(docx_file)
    print(f"✓ Saved to {docx_file}")


def main():
    """Main conversion function."""
    base_dir = Path(__file__).parent

    # Convert Production Readiness document
    production_md = base_dir / "PRODUCTION_READINESS.md"
    production_docx = base_dir / "PRODUCTION_READINESS.docx"
    if production_md.exists():
        convert_markdown_to_docx(production_md, production_docx)
    else:
        print(f"Warning: {production_md} not found")

    # Convert Codebase Documentation
    codebase_md = base_dir / "CODEBASE_DOCUMENTATION.md"
    codebase_docx = base_dir / "CODEBASE_DOCUMENTATION.docx"
    if codebase_md.exists():
        convert_markdown_to_docx(codebase_md, codebase_docx)
    else:
        print(f"Warning: {codebase_md} not found")

    print("\n✓ Conversion complete!")
    print(f"  - {production_docx}")
    print(f"  - {codebase_docx}")


if __name__ == "__main__":
    main()
